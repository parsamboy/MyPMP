# -*- coding: utf-8 -*-
"""گزارش مقایسه نسخه Up-v2 با نسخه قبلی."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "گزارش مقایسه نسخه جدید.docx"
FA, TITR, LAT = "B Lotus", "B Titr", "Times New Roman"
NAVY=(0x1F,0x38,0x64); RED=(0xC0,0x00,0x00); GREEN=(0x1E,0x6B,0x2E)
GREY=(0x59,0x59,0x59); ORANGE=(0xB0,0x5A,0x00)


def set_rtl(p):
    b=OxmlElement("w:bidi"); b.set(qn("w:val"),"1"); p._p.get_or_add_pPr().append(b)

def sr(r,size=12,bold=False,color=None,latin=False,font=None):
    f=font or (LAT if latin else FA)
    r.font.name=f; r.font.size=Pt(size); r.font.bold=bold
    if color: r.font.color.rgb=RGBColor(*color)
    rPr=r._r.get_or_add_rPr(); rf=rPr.find(qn("w:rFonts"))
    if rf is None: rf=OxmlElement("w:rFonts"); rPr.append(rf)
    for a in ("w:cs","w:ascii","w:hAnsi"): rf.set(qn(a),f)
    s=OxmlElement("w:szCs"); s.set(qn("w:val"),str(int(size*2))); rPr.append(s)
    if not latin:
        rt=OxmlElement("w:rtl"); rt.set(qn("w:val"),"1"); rPr.append(rt)
    lg=OxmlElement("w:lang"); lg.set(qn("w:bidi"),"fa-IR"); rPr.append(lg)

def para(doc,text,size=12,bold=False,align="just",color=None,sb=0,sa=4,font=None,indent=None):
    p=doc.add_paragraph(); set_rtl(p)
    p.alignment={"just":WD_ALIGN_PARAGRAPH.JUSTIFY,"right":WD_ALIGN_PARAGRAPH.RIGHT,
                 "center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    pf=p.paragraph_format; pf.space_before=Pt(sb); pf.space_after=Pt(sa); pf.line_spacing=1.15
    if indent is not None: pf.right_indent=Cm(indent)
    if text: sr(p.add_run(text),size,bold,color,font=font)
    return p

def shade(c,hx):
    sh=OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hx)
    c._tc.get_or_add_tcPr().append(sh)

def ct(cell,text,size=10,bold=False,color=None,align="right",latin=False):
    cell.text=""; p=cell.paragraphs[0]; set_rtl(p)
    p.alignment={"right":WD_ALIGN_PARAGRAPH.RIGHT,"center":WD_ALIGN_PARAGRAPH.CENTER,
                 "just":WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    p.paragraph_format.space_before=Pt(1.5); p.paragraph_format.space_after=Pt(1.5)
    for i,l in enumerate(str(text).split("\n")):
        if i:
            p=cell.add_paragraph(); set_rtl(p); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before=Pt(1.5); p.paragraph_format.space_after=Pt(1.5)
        sr(p.add_run(l),size,bold,color,latin=latin)

def borders(t):
    b=OxmlElement("w:tblBorders")
    for e in ("top","left","bottom","right","insideH","insideV"):
        x=OxmlElement("w:"+e); x.set(qn("w:val"),"single"); x.set(qn("w:sz"),"6")
        x.set(qn("w:color"),"808080"); b.append(x)
    t._tbl.tblPr.append(b); t._tbl.tblPr.append(OxmlElement("w:bidiVisual"))

def table(doc,hdr,rows,widths=None,fill="DCE6F1",latin_cols=(),sizes=10):
    t=doc.add_table(rows=1+len(rows),cols=len(hdr)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    borders(t)
    for i,h in enumerate(hdr):
        ct(t.cell(0,i),h,size=sizes,bold=True,align="center"); shade(t.cell(0,i),fill)
    for ri,row in enumerate(rows,1):
        for ci,v in enumerate(row):
            ct(t.cell(ri,ci),v,size=sizes,align="center" if ci==0 else "right",
               latin=ci in latin_cols)
            if ri%2==1: shade(t.cell(ri,ci),"F7F9FC")
    if widths:
        for i,w in enumerate(widths):
            for r in t.rows: r.cells[i].width=Cm(w)
    return t


FIXED = [
 ("پانویس‌ها", "۸۵ → ۹۸ پانویس", "۱۳ پانویس تازه اضافه شد؛ از جمله Ryff، Ryan & Deci، Diener و Salkovskis که در گزارش قبلی به عنوان «جاافتاده» فهرست کرده بودم."),
 ("منابع", "۸۴ → ۱۲۰ مدخل", "۳۶ مدخل خالص افزوده شد (لاتین از ۷۵ به ۱۰۴). بسیاری از منابعی که در متن استناد شده ولی در فهرست نبودند، اکنون اضافه شده‌اند."),
 ("جدول ۴-۸ — مقدار F", "هر چهار گام اصلاح شد", "۴۱/۸۷ → ۴۱/۸ ، ۴۴/۹۵ → ۳۱/۲۷ ، ۴۳/۱۲ → ۲۴/۱۵ ، ۴۵/۳۸ → ۲۰/۸۲. اکنون هر چهار مقدار دقیقاً با MS رگرسیون ÷ MS باقیمانده می‌خوانند. این مهم‌ترین اصلاح علمی نسخه جدید است."),
 ("ارجاع جدول ۴-۶", "«جدول ۴-۹» → «جدول ۴-۶»", "ارجاع متقابل غلط در فرضیه فرعی اول اصلاح شد."),
 ("ارجاع جدول ۴-۷", "«جدول ۴-۶» → «جدول ۴-۷»", "ارجاع متقابل غلط در فرضیه فرعی دوم اصلاح شد."),
 ("ارجاع جدول ۴-۸", "«جدول ۴-۱۲» → «جدول ۴-۸»", "ارجاع به جدولی که وجود نداشت اصلاح شد."),
 ("جدول ۴-۴", "«sd» → «انحراف معیار»", "نماد لاتین با معادل فارسی جایگزین شد."),
 ("پانویس ۵۸ / ۶۵", "American Psychiatric Association", "ترتیب معکوس کلمات در یکی از دو نمونه اصلاح شد."),
 ("عنوان جدول ۴-۷", "«اضطراب سلامت کل» → «اضطراب سلامتی کل»", "یکدست‌سازی اصطلاح."),
]

REMAIN = [
 ("۱", "درجه آزادی باقیمانده", "۷۹، ۷۸، ۷۷، ۷۶", "۷۸، ۷۷، ۷۶، ۷۵",
  "با n=۸۰ باید n−k−۱ باشد. این تنها ایراد «بحرانی» است که دست‌نخورده مانده."),
 ("۲", "سقف مقیاس اضطراب مرگ", "بیشینه ۷۰ ، میانگین ۴۸/۲۳", "دامنه ۰ تا ۱۵",
  "فصل سوم همچنان می‌گوید تمپلر ۱۵ ماده بلی/خیر است. تناقض پابرجاست."),
 ("۳", "سقف مقیاس اضطراب سلامتی", "بیشینه ۶۸", "حداکثر ۵۴ (۱۸×۳)",
  "بدون تغییر."),
 ("۴", "مجموع مجذورات گام ۳", "۸۹۸۹/۵۷", "۹۰۱۶/۵۷",
  "همچنان ۲۷ واحد با سه گام دیگر اختلاف دارد. SS باقیمانده باید ۴۶۵۸/۳۶ باشد."),
 ("۵", "ساختار جدول ۴-۳", "«زیر دیپلم/دیپلم» در یک سلول", "دو ردیف مستقل",
  "بدون تغییر."),
 ("۶", "زیرنویس ستاره‌ها", "ندارد", "** p<۰/۰۱ ، n=۸۰",
  "جدول‌های ۴-۶ و ۴-۷ همچنان بدون یادداشت هستند."),
 ("۷", "هم‌خطی (VIF)", "گزارش نشده", "افزودن جمله VIF",
  "نمره کل هوش معنوی همچنان کنار دو خرده‌مقیاس خودش در مدل است."),
 ("۸", "جدول K–S", "بدون عدد", "افزودن مقادیر p",
  "بدون تغییر."),
]

NEWBUGS = [
 ("خطای نو ۱", "بحرانی", "متن تفسیری با جدول ۴-۸ نمی‌خواند",
  "R² گام ۳ از ۵۰/۰ به ۴۸/۰ و گام ۴ از ۵۴/۰ به ۵۲/۰ تغییر کرد، اما متن زیر جدول هنوز می‌گوید "
  "«مقدار ۵۴/۰=R۲ … ۵۴ درصد» و سهم‌ها را «۳۴ + ۱۱ + ۵ + ۴» می‌شمارد. "
  "با اعداد جدید مجموع ۵۲ درصد و سهم گام ۳ برابر ۳ درصد است.",
  "یا متن را به ۵۲ درصد و «۳۴ + ۱۱ + ۳ + ۴» اصلاح کنید، یا R² جدول را به مقادیر قبلی بازگردانید. "
  "توصیه من: جدول بازسازی‌شده گزارش قبلی را یکجا جایگزین کنید تا همه‌چیز هم‌خوان شود."),
 ("خطای نو ۲", "مهم", "ردیف «ایجاد معنای شخصی» در جدول ۴-۹ جابه‌جا شده",
  "در فایل جدید این ردیف به صورت «۲۴/۰-  [سه سلول خالی]  ۰۹/۰  ۲۵/۰-  ۶۷/۲-  ۰۰۹/۰» ذخیره شده "
  "است، یعنی یک سلول خالی اضافه وارد شده و ستون‌ها یک واحد لغزیده‌اند.",
  "ردیف باید دقیقاً چنین باشد: B=۲۴/۰- ، خطای استاندارد=۰۹/۰ ، β=۲۵/۰- ، t=۶۷/۲- ، معناداری=۰۰۹/۰"),
 ("خطای نو ۳", "متوسط", "پانویس ۲۱: «Rajer Gould»",
  "املای درست نام کوچک Roger است، نه Rajer. ضمناً طبق قاعده شما پانویس باید فقط نام خانوادگی باشد.",
  "به «Gould» تغییر یابد."),
 ("خطای نو ۴", "متوسط", "پانویس ۷۲: «Gouldberg»",
  "این املا وجود ندارد. با توجه به متن (انجمن روان‌شناسی آمریکا، درمان شناختی-رفتاری) "
  "احتمالاً Goldberg درست است.",
  "به «Goldberg» تغییر یابد و از Gould (پانویس ۷۴) تفکیک شود."),
 ("خطای نو ۵", "جزئی", "سه مدخل تکراری در فهرست منابع",
  "Atchley (۲۰۱۶)، Harding (۲۰۰۵) و Noyes (۲۰۰۵) هر کدام دوبار در فهرست منابع آمده‌اند.",
  "نسخه تکراری هر سه حذف شود."),
 ("خطای نو ۶", "جزئی", "جمله توصیفی جدول ۴-۵ گسترده‌تر شد",
  "جمله به «میانگین هوش معنوی و اضطراب سلامتی … در سطح نسبتاً بالا» تغییر کرده است. "
  "با توجه به آنکه بیشینه اضطراب سلامتی (۶۸) از سقف مقیاس (۵۴) بیشتر است، "
  "این داوری اکنون ادعای بیشتری می‌کند که پشتوانه ندارد.",
  "تا زمان روشن شدن دامنه واقعی مقیاس، این توصیف کیفی حذف یا مشروط شود."),
]

STILL_MISSING_FN = [
 ("تمپلر", "Templer", "سازنده ابزار اصلی؛ همچنان بدون پانویس و بدون مدخل در منابع"),
 ("یونگ", "Jung", "۵ بار در متن؛ پانویس Wagenseller اضافه شد ولی خودِ یونگ پانویس ندارد"),
 ("فرانکل", "Frankl", "یک نظریه است؛ بدون پانویس"),
 ("آدلر", "Adler", "ارجاع به «نظریه قدرت‌طلبی آدلر»؛ بدون پانویس"),
 ("ویلیام جیمز", "William James", "بدون پانویس"),
 ("کیز", "Keyes", "منبع Keyes (2002) به فهرست اضافه شد ولی پانویس در متن نه"),
 ("فرنس", "؟", "همچنان نامشخص؛ احتمالاً Franz"),
 ("نظریه مدیریت وحشت", "Terror Management Theory", "عنوان بخش ۲-۲-۲-۳؛ فقط مخفف TMT آمده"),
]


def build():
    doc=Document(); s=doc.sections[0]
    s.page_width,s.page_height=Cm(21),Cm(29.7)
    s.top_margin=s.bottom_margin=Cm(2.0); s.left_margin=s.right_margin=Cm(1.8)
    b=OxmlElement("w:bidi"); b.set(qn("w:val"),"1"); s._sectPr.append(b)
    doc.styles["Normal"].font.name=FA; doc.styles["Normal"].font.size=Pt(12)

    para(doc,"گزارش مقایسه نسخه جدید",20,True,"center",NAVY,font=TITR,sa=2)
    para(doc,"مقایسه (Up-v2) با نسخه پیشین — پانویس‌ها، منابع و محتوای علمی فصل چهارم",
         12,False,"center",GREY,sa=10)

    para(doc,"جمع‌بندی",14,True,"right",NAVY,font=TITR,sa=4)
    para(doc,"پیشرفت نسخه جدید چشمگیر است. مهم‌ترین دستاورد آن اصلاح مقادیر F در جدول ۴-۸ است "
             "که در نسخه قبلی با هیچ ترکیبی از اعداد جدول بازتولید نمی‌شد و اکنون هر چهار گام "
             "دقیقاً درست محاسبه شده است. فهرست منابع نیز از ۸۴ به ۱۲۰ مدخل رسیده و ۱۳ پانویس تازه افزوده شده است.",
         12,sa=5)
    para(doc,"در عین حال، ۸ ایراد از گزارش قبلی هنوز اصلاح نشده و ۶ ایراد تازه در جریان همین "
             "ویرایش پدید آمده است — که یکی از آن‌ها بحرانی است: متن تفسیری دیگر با جدول اصلاح‌شده نمی‌خواند.",
         12,color=ORANGE,sa=8)

    table(doc,["شاخص","نسخه قبلی","نسخه جدید"],
          [["تعداد پانویس","۸۵","۹۸"],["مدخل‌های فهرست منابع","۸۴","۱۲۰"],
           ["منابع لاتین","۷۵","۱۰۴"],["حجم متن اصلی","۱۲۶٬۳۰۴ نویسه","۱۳۶٬۴۱۳ نویسه"],
           ["مقادیر F سازگار در جدول ۴-۸","۱ از ۴","۴ از ۴"],
           ["پانویس دارای سال","۰","۰ (منطبق با قاعده)"]],
          widths=[7.0,4.8,4.8])

    doc.add_page_break()
    para(doc,"۱) آنچه اصلاح شده است",15,True,"right",GREEN,font=TITR,sa=4)
    table(doc,["مورد","تغییر","توضیح"],
          [[a,b_,c] for a,b_,c in FIXED],widths=[3.4,4.0,9.2],fill="E2EFDA")

    doc.add_page_break()
    para(doc,"۲) ایرادهای گزارش قبلی که هنوز باقی است",15,True,"right",RED,font=TITR,sa=4)
    table(doc,["ش","مورد","وضعیت فعلی","مقدار درست","توضیح"],
          [[a,b_,c,d,e] for a,b_,c,d,e in REMAIN],
          widths=[1.0,3.4,3.6,3.4,5.6],fill="FBE4E4")

    doc.add_page_break()
    para(doc,"۳) خطاهای تازه‌ای که در این ویرایش پدید آمده",15,True,"right",ORANGE,
         font=TITR,sa=4)
    para(doc,"این موارد در نسخه قبلی وجود نداشتند و حاصل خودِ ویرایش هستند:",
         11,color=GREY,sa=5)
    for code,sev,title,what,fix in NEWBUGS:
        t=doc.add_table(rows=3,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; borders(t)
        col=RED if sev=="بحرانی" else (ORANGE if sev=="مهم" else GREY)
        fillc="FBE4E4" if sev=="بحرانی" else ("FDF0E0" if sev=="مهم" else "F0F0F0")
        ct(t.cell(0,0),f"{code}  ·  {sev}  ·  {title}",size=11.5,bold=True,color=col)
        shade(t.cell(0,0),fillc)
        ct(t.cell(1,0),what,size=11,align="just")
        ct(t.cell(2,0),"اصلاح: "+fix,size=11,align="just",color=GREEN)
        shade(t.cell(2,0),"EDF5EE")
        para(doc,"",size=6,sa=6)

    doc.add_page_break()
    para(doc,"۴) پانویس‌هایی که همچنان جا افتاده‌اند",15,True,"right",NAVY,font=TITR,sa=4)
    para(doc,"از ۱۵ موردی که در گزارش قبلی فهرست کردم، ۵ مورد (Ryff، Ryan، Deci، Diener، Salkovskis) "
             "اضافه شده و ۸ مورد زیر همچنان باقی است:",11,color=GREY,sa=5)
    table(doc,["واژه در متن","پانویس پیشنهادی","توضیح"],
          [[a,b_,c] for a,b_,c in STILL_MISSING_FN],
          widths=[3.2,4.0,9.4],fill="FBE4E4",latin_cols=(1,))

    para(doc,"۵) پانویس‌های با املای غلط که اصلاح نشده‌اند",15,True,"right",ORANGE,
         font=TITR,sb=10,sa=4)
    table(doc,["ش","متن فعلی","املای درست"],
          [["۱۳","1 Gerontology","Gerontology"],
           ["۱۶","Nikolich","Nikolich-Žugich"],
           ["۱۷","Fulop","Fülöp"],
           ["۲۲","3 Levinson","Levinson"],
           ["۲۳","4 SHAie","Schaie"],
           ["۲۴","5 Erikson","Erikson"],
           ["۳۹","Greenberg","Greenberg, Pyszczynski & Solomon"],
           ["۷۸","Wang &Zhao","Wang & Zhao"],
           ["۸۰","Yonker, Schanbelrauch & Dehaan","Yonker, Schnabelrauch & DeHaan"],
           ["۸۳","Suad,Zarina Mat, Zulkarnin A,Hatta Nuria","Mat Saad, Hatta & Mohamad"],
           ["۸۶","gin & Purohit","Jain & Purohit"],
           ["۸۸","Dorson & Poul","Thorson & Powell"],
           ["۸۹","Moreiva & Almeida","Moreira-Almeida & Koenig"]],
          widths=[1.4,7.0,8.2],fill="FDF0E0",latin_cols=(1,2))
    para(doc,"نکته: پانویس‌های ۱۳، ۲۲، ۲۳ و ۲۴ هنوز عدد چسبیده به متن دارند. این اعداد بازمانده "
             "شماره‌گذاری دستی هستند و باید حذف شوند.",11,color=GREY,sb=4,sa=6)

    doc.add_page_break()
    para(doc,"۶) منابعی که همچنان در فهرست نیستند",15,True,"right",NAVY,font=TITR,sa=4)
    para(doc,"این آثار در متن استناد شده‌اند ولی مدخل ندارند:",11,color=GREY,sa=5)
    table(doc,["اثر","محل استناد در متن"],
          [["Templer, D. I. (1970)","فصل سوم — ابزار اصلی سنجش اضطراب مرگ"],
           ["Krejcie, R. V., & Morgan, D. W. (1970)","فصل سوم — تعیین حجم نمونه"],
           ["Thorson & Powell (1988)","پیشینه خارجی (پانویس ۸۸)"],
           ["Levinson, D. J. (1978)","نظریه‌های روان‌شناختی سالمندی"],
           ["Schaie, K. W. (1979)","نظریه‌های روان‌شناختی سالمندی"],
           ["Greenberg, Pyszczynski & Solomon (1986)","نظریه مدیریت وحشت"],
           ["Mikulincer & Florian (2000)","نظریه دلبستگی"],
           ["Frankl, V. E.","هوش معنوی — معناجویی"],
           ["Papalia & Martorell","سالمندی"]],
          widths=[6.4,10.2],fill="FBE4E4",latin_cols=(0,))

    para(doc,"اولویت پیشنهادی",15,True,"right",NAVY,font=TITR,sb=10,sa=4)
    for i,t in enumerate([
      "«خطای نو ۱» را فوراً حل کنید — متن و جدول ۴-۸ اکنون با هم در تناقض‌اند و این بدتر از حالت قبلی است.",
      "ردیف جابه‌جاشده جدول ۴-۹ را درست کنید (خطای نو ۲).",
      "درجه آزادی‌ها را به ۷۸/۷۷/۷۶/۷۵ اصلاح کنید و SS گام ۳ را به ۴۶۵۸/۳۶ ببرید.",
      "مدخل Templer و Krejcie & Morgan را به منابع بیفزایید — نبود آن‌ها در جلسه دفاع پرسیده می‌شود.",
      "۱۳ پانویس با املای غلط و ۴ عدد چسبیده را اصلاح کنید.",
      "سه مدخل تکراری منابع را حذف کنید.",
    ],1):
        para(doc,f"{'۱۲۳۴۵۶'[i-1]}. "+t,11.5,sa=5,indent=0.4)

    doc.save(OUT); print("saved",OUT,os.path.getsize(OUT))


if __name__=="__main__":
    build()
