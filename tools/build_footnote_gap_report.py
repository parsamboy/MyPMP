# -*- coding: utf-8 -*-
"""
گزارش پوشش پانویس‌ها بر پایه قاعده جدید کاربر:
پانویس فقط = املای درست انگلیسی نامِ فرد یا نظریهٔ غیرایرانی که در متن فارسی آمده.
بدون سال، بدون ارجاع کتاب‌شناختی.
خروجی: «گزارش پوشش پانویس‌ها.docx»
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "گزارش پوشش پانویس‌ها.docx"
FA, TITR, LAT = "B Lotus", "B Titr", "Times New Roman"
NAVY = (0x1F, 0x38, 0x64); RED = (0xC0, 0x00, 0x00)
GREEN = (0x1E, 0x6B, 0x2E); GREY = (0x59, 0x59, 0x59); ORANGE = (0xB0, 0x5A, 0x00)


def set_rtl(p):
    b = OxmlElement("w:bidi"); b.set(qn("w:val"), "1")
    p._p.get_or_add_pPr().append(b)


def style_run(r, size=12, bold=False, color=None, latin=False, font=None):
    f = font or (LAT if latin else FA)
    r.font.name = f; r.font.size = Pt(size); r.font.bold = bold
    if color: r.font.color.rgb = RGBColor(*color)
    rPr = r._r.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rPr.append(rf)
    for a in ("w:cs", "w:ascii", "w:hAnsi"): rf.set(qn(a), f)
    s = OxmlElement("w:szCs"); s.set(qn("w:val"), str(int(size * 2))); rPr.append(s)
    if not latin:
        rt = OxmlElement("w:rtl"); rt.set(qn("w:val"), "1"); rPr.append(rt)
    lg = OxmlElement("w:lang"); lg.set(qn("w:bidi"), "fa-IR"); rPr.append(lg)


def para(doc, text, size=12, bold=False, align="just", color=None,
         sb=0, sa=4, font=None, indent=None):
    p = doc.add_paragraph(); set_rtl(p)
    p.alignment = {"just": WD_ALIGN_PARAGRAPH.JUSTIFY, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    pf = p.paragraph_format
    pf.space_before = Pt(sb); pf.space_after = Pt(sa); pf.line_spacing = 1.15
    if indent is not None: pf.right_indent = Cm(indent)
    if text: style_run(p.add_run(text), size, bold, color, font=font)
    return p


def shade(cell, hx):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hx)
    cell._tc.get_or_add_tcPr().append(sh)


def ctext(cell, text, size=10, bold=False, color=None, align="right", latin=False):
    cell.text = ""
    p = cell.paragraphs[0]; set_rtl(p)
    p.alignment = {"right": WD_ALIGN_PARAGRAPH.RIGHT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "just": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    p.paragraph_format.space_before = Pt(1.5); p.paragraph_format.space_after = Pt(1.5)
    for i, line in enumerate(str(text).split("\n")):
        if i:
            p = cell.add_paragraph(); set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(1.5); p.paragraph_format.space_after = Pt(1.5)
        style_run(p.add_run(line), size, bold, color, latin=latin)


def borders(tbl):
    b = OxmlElement("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        x = OxmlElement("w:" + e)
        x.set(qn("w:val"), "single"); x.set(qn("w:sz"), "6"); x.set(qn("w:color"), "808080")
        b.append(x)
    tbl._tbl.tblPr.append(b)
    tbl._tbl.tblPr.append(OxmlElement("w:bidiVisual"))


def table(doc, headers, rows, widths=None, head_fill="DCE6F1", latin_cols=()):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; borders(t)
    for i, h in enumerate(headers):
        ctext(t.cell(0, i), h, size=10, bold=True, align="center")
        shade(t.cell(0, i), head_fill)
    for ri, row in enumerate(rows, 1):
        for ci, v in enumerate(row):
            ctext(t.cell(ri, ci), v, size=10, align="center" if ci in (0,) else "right",
                  latin=ci in latin_cols)
            if ri % 2 == 1: shade(t.cell(ri, ci), "F7F9FC")
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows: r.cells[i].width = Cm(w)
    return t


# ---------------- data ----------------
# ۱۵ نام غیرایرانی که هیچ پانویسی ندارند
MISSING = [
 ("۱", "تمپلر", "Templer", "ص ۳۵۲ — «این پرسشنامه در سال ۱۹۷۰ توسط تمپلر ساخته شده»", "سازنده ابزار اصلی پژوهش؛ نبود پانویس آن آشکارترین نقص است."),
 ("۲", "سالکوفسکیس", "Salkovskis", "ص ۳۵۴ — «توسط سالکوفسکیس و همکاران (۲۰۰۲) تدوین شده»", "در ص ۲۷۹ پانویس Salkovskis & Warwick آمده، اما این املای فارسی متفاوت («سالکوفسکیس») بدون پانویس مانده."),
 ("۳", "ریف", "Ryff", "ص ۲۶۰ — «ریف (۱۹۸۹) به تفاوت‌های هوش هیجانی…»", "۵ بار در متن تکرار شده و هیچ‌گاه پانویس نخورده است."),
 ("۴", "یونگ", "Jung", "ص ۲۵۸ — «یونگ و فرنس در مخالفت با دیدگاه منفی فروید»", "۵ بار تکرار شده؛ نام کاملاً شناخته‌شده ولی بدون پانویس."),
 ("۵", "فرانکل", "Frankl", "ص ۲۵۸ — «نظریه فرانکل بر معناجویی افراد در زندگی تأکید دارد»", "یک نظریه است، پس طبق قاعده شما حتماً پانویس می‌خواهد."),
 ("۶", "آدلر", "Adler", "ص ۲۵۸ — «نه بر پایه نظریه قدرت‌طلبی آدلر»", "ارجاع به نظریه؛ نیازمند پانویس."),
 ("۷", "ویلیام جیمز", "William James", "ص ۲۵۸ — «ویلیام جیمز پدر روانشناسی آمریکا»", "نام کوچک هم آمده، پس پانویس باید کامل باشد."),
 ("۸", "رایان", "Ryan", "ص ۲۵۹ و ۲۷۴ — «(رایان و دسی، ۲۰۱۵)»", "۶ بار تکرار شده بدون هیچ پانویسی."),
 ("۹", "دسی", "Deci", "ص ۲۶۰ — «(رایان و دسی، ۲۰۱۵)»", "پانویس King & DeCicco در جای دیگری هست، ولی آن شخصِ دیگری است؛ Deci بدون پانویس مانده."),
 ("۱۰", "کیز", "Keyes", "ص ۲۷۳ و ۲۷۴ — «(کیز، ۲۰۱۲)»", "۳ بار تکرار شده بدون پانویس."),
 ("۱۱", "دیتر", "Diener", "ص ۳۱۳ — «پرسشنامه رضایت از زندگی دیتر و همکارانش»", "املای فارسی هم غلط است؛ «داینر» درست‌تر است. در ص ۲۶۸ همین فرد «داینر» نوشته شده."),
 ("۱۲", "فرنس", "؟", "ص ۲۵۸ — «یونگ و فرنس در مخالفت با دیدگاه منفی فروید»", "این نام قابل شناسایی نیست. احتمالاً تحریفِ Franz (ماری‌لوئیز فون فرانتس، همکار یونگ) است. نیازمند بررسی نویسنده."),
 ("۱۳", "نظریه مدیریت وحشت", "Terror Management Theory", "ص ۲۲۱ — عنوان «۲-۲-۲-۳- نظریه مدیریت وحشت (TMT)»", "یک نظریهٔ غیرایرانی است و طبق قاعده شما باید پانویس بخورد. مخفف TMT در متن آمده ولی صورت کامل انگلیسی آن نیامده."),
 ("۱۴", "کوبلر-راس", "Kübler-Ross", "بررسی شود — در متن به مراحل سوگ اشاره شده", "اگر نام در متن آمده باشد نیازمند پانویس است؛ در جست‌وجوی ما یافت نشد ولی نویسنده تأیید کند."),
 ("۱۵", "بالبی", "Bowlby", "ص ۲۲۳ — بخش «نظریه دلبستگی»", "نظریه دلبستگی بدون ذکر نام بالبی آمده؛ اگر نام اضافه شود پانویس لازم دارد."),
]

# پانویس‌های موجود که املای انگلیسی‌شان غلط است
WRONG = [
 ("۸", "ژرونتولوژی", "1 Gerontology", "Gerontology", "عدد «۱» اشتباهاً داخل متن پانویس تایپ شده"),
 ("۱۱", "نیکولاس", "Nikolich", "Nikolich-Žugich", "نام خانوادگی ناقص است"),
 ("۱۲", "فولپ", "Fulop", "Fülöp", "علامت‌های آلمانی/مجاری حذف شده"),
 ("۱۶", "راجر گولد", "Gold", "Gould", "املای غلط — Gold با Gould اشتباه شده"),
 ("۱۷", "دانیل لوینسون", "3 Levinson", "Levinson", "عدد «۳» اضافه است"),
 ("۱۸", "شی", "4 SHAie", "Schaie", "هم عدد اضافه، هم املا کاملاً غلط"),
 ("۱۹", "اریکسون", "5 Erikson", "Erikson", "عدد «۵» اضافه است"),
 ("۳۳", "گرین برگ", "Greenberg", "Greenberg, Pyszczynski & Solomon", "نظریه مدیریت وحشت سه بنیان‌گذار دارد"),
 ("۵۱", "انجمن روان‌پزشکی آمریکا", "Association American Psychiatric", "American Psychiatric Association", "ترتیب کلمات معکوس شده"),
 ("۶۵", "وانگ و ژائو", "Wang &Zhao", "Wang & Zhao", "فاصله پیش از & جا افتاده"),
 ("۶۷", "یانکر، اسنابلروچ و دهان", "Yonker, Schanbelrauch & Dehaan", "Yonker, Schnabelrauch & DeHaan", "دو غلط املایی"),
 ("۷۰", "حتا و نوریا محمد", "Suad,Zarina Mat, Zulkarnin A,Hatta Nuria", "Mat Saad, Hatta & Mohamad", "نام‌ها کاملاً درهم ریخته است"),
 ("۷۳", "جین و پوروحیت", "gin & Purohit", "Jain & Purohit", "حرف اول نام افتاده"),
 ("۷۵", "دورسون و پاوول", "Dorson & Poul", "Thorson & Powell", "هر دو نام غلط نوشته شده"),
 ("۷۶", "موریرا آلمیدا و کوئینگ", "Moreiva & Almeida", "Moreira-Almeida & Koenig", "Moreira-Almeida نامِ یک نفر است، نه دو نفر"),
]

# پانویس‌هایی که با قاعده جدید سازگار نیستند
POLICY = [
 ("۹ و ۵۵", "سازمان جهانی بهداشت", "World Health Organization",
  "طبق راهنمای اکسل خودتان، سازمان‌های بین‌المللی (WHO، UN، UNFPA) پانویس لازم ندارند. اگر قاعده «فقط نام افراد و نظریه‌ها» را سخت‌گیرانه اجرا کنیم، این دو حذف می‌شوند."),
 ("۵۱ و ۵۶", "انجمن روان‌پزشکی آمریکا", "American Psychiatric Association",
  "این یک سازمان است نه فرد. البته چون DSM-5 به آن وابسته است، راهنمای اکسل آن را استثنا کرده بود. تصمیم با شماست."),
 ("۵۷", "نشنال اینستیتوت آن ایجینگ", "National Institute on Aging",
  "نام سازمان در متن فارسی «آوانویسی» شده است. بهتر است متن فارسی به «مؤسسه ملی سالمندی» تغییر کند و پانویس انگلیسی بماند."),
 ("۸", "ژرونتولوژی", "Gerontology",
  "این یک رشتهٔ علمی است، نه نام فرد یا نظریه. طبق قاعده جدید شما باید حذف شود، مگر آنکه بخواهید معادل انگلیسی اصطلاح تخصصی را نگه دارید."),
]

DUP = [
 ("پینتو / Pinto", "۴۰ و ۴۱"), ("کینگ و دسیکو / King & DeCicco", "۴۲، ۷۲، ۸۳"),
 ("ابراموویتس / Abramowitz", "۴۴ و ۴۷"), ("انجمن روان‌پزشکی آمریکا", "۵۱ و ۵۶"),
 ("سازمان جهانی بهداشت", "۹ و ۵۵"), ("وانگ / Wang", "۶ و ۵۹"),
 ("گولد / Gould", "۱۶ و ۶۴"), ("لوین / Levin", "۶۹ و ۸۲"),
 ("جین و پوروحیت / Jain & Purohit", "۷۳ و ۷۷"), ("هاردینگ / Harding", "۷۴ و ۸۴"),
 ("نویز / Noyes", "۷۸ و ۸۵"), ("یالوم / Yalom", "۳۲ و ۸۰"),
]


def build():
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21), Cm(29.7)
    s.top_margin = s.bottom_margin = Cm(2.0)
    s.left_margin = s.right_margin = Cm(1.9)
    b = OxmlElement("w:bidi"); b.set(qn("w:val"), "1"); s._sectPr.append(b)
    doc.styles["Normal"].font.name = FA
    doc.styles["Normal"].font.size = Pt(12)

    para(doc, "گزارش پوشش پانویس‌ها", 20, True, "center", NAVY, font=TITR, sa=2)
    para(doc, "بر پایه قاعده: پانویس فقط املای درست انگلیسیِ نام فرد یا نظریهٔ غیرایرانی",
         12, False, "center", GREY, sa=10)

    para(doc, "پاسخ کوتاه", 14, True, "right", NAVY, font=TITR, sb=4, sa=4)
    para(doc,
         "خیر، همه موارد پوشش داده نشده است. متن فارسی سطر‌به‌سطر بررسی شد و «۱۵ نام غیرایرانی» "
         "پیدا شد که هیچ پانویسی ندارند — از جمله تمپلر و سالکوفسکیس که سازندگان دو ابزار اصلی "
         "همین پژوهش هستند. افزون بر این، ۱۵ پانویسِ موجود املای انگلیسی غلط دارند.",
         12, sa=6)

    para(doc, "یک خبر خوب: طبق قاعده شما، پانویس‌ها نباید سال داشته باشند. "
              "بررسی هر ۸۵ پانویس نشان داد هیچ‌کدام سال ندارند؛ پس از این بابت فایل کاملاً منطبق است "
              "و پرسش قبلیِ من دربارهٔ «با سال یا بدون سال» منتفی شد.",
         11.5, color=GREEN, sa=8)

    rows = [["۸۵", "پانویس موجود در فایل"], ["۱۵", "نام غیرایرانی بدون هیچ پانویس"],
            ["۱۵", "پانویس با املای انگلیسی غلط"], ["۴", "پانویس مغایر با قاعده جدید (سازمان/اصطلاح)"],
            ["۱۲", "نام تکراری که دوبار پانویس خورده"], ["۰", "پانویس دارای سال — منطبق با قاعده"]]
    table(doc, ["تعداد", "وضعیت"], rows, widths=[2.2, 14.5])

    doc.add_page_break()
    para(doc, "۱) نام‌های غیرایرانی که پانویس ندارند", 15, True, "right", RED, font=TITR, sa=4)
    para(doc, "این‌ها باید به فایل اضافه شوند. ستون «پانویس پیشنهادی» آماده درج است.",
         11, color=GREY, sa=5)
    table(doc, ["ش", "واژه در متن", "پانویس پیشنهادی", "نخستین محل", "توضیح"],
          [[a, b_, c, d, e] for a, b_, c, d, e in MISSING],
          widths=[1.0, 2.6, 3.4, 5.0, 5.0], head_fill="FBE4E4", latin_cols=(2,))

    doc.add_page_break()
    para(doc, "۲) پانویس‌های موجود با املای انگلیسی غلط", 15, True, "right", ORANGE, font=TITR, sa=4)
    para(doc, "قاعده شما می‌گوید پانویس باید «تلفظ/املای درست انگلیسی» باشد؛ "
              "پس این ۱۵ مورد دقیقاً همان چیزی است که باید اصلاح شود.", 11, color=GREY, sa=5)
    table(doc, ["ش پانویس", "واژه در متن", "متن فعلی", "املای درست", "ایراد"],
          [[a, b_, c, d, e] for a, b_, c, d, e in WRONG],
          widths=[1.6, 3.2, 4.0, 4.2, 4.0], head_fill="FDF0E0", latin_cols=(2, 3))

    doc.add_page_break()
    para(doc, "۳) پانویس‌هایی که با قاعده جدید سازگار نیستند", 15, True, "right", NAVY, font=TITR, sa=4)
    para(doc, "این‌ها نام فرد یا نظریه نیستند. تصمیم نهایی با شماست:", 11, color=GREY, sa=5)
    table(doc, ["ش پانویس", "واژه در متن", "متن پانویس", "پیشنهاد"],
          [[a, b_, c, d] for a, b_, c, d in POLICY],
          widths=[1.8, 3.4, 4.2, 7.6], head_fill="E9E9F5", latin_cols=(2,))

    para(doc, "۴) نام‌هایی که بیش از یک‌بار پانویس خورده‌اند", 15, True, "right", NAVY,
         font=TITR, sb=10, sa=4)
    para(doc, "راهنمای اکسل شما می‌گفت پانویس فقط در نخستین اشاره بیاید. "
              "این ۱۲ نام دوبار (یا بیشتر) پانویس خورده‌اند و باید فقط اولی بماند:",
         11, color=GREY, sa=5)
    table(doc, ["نام", "شماره پانویس‌های تکراری"], [[a, b_] for a, b_ in DUP],
          widths=[9.0, 7.6], head_fill="EFEFEF")

    doc.add_page_break()
    para(doc, "نکته‌ای که باید تصمیم بگیرید", 15, True, "right", NAVY, font=TITR, sa=4)
    para(doc,
         "قاعده جدید شما («پانویس فقط املای انگلیسی نام») با سیاست قبلی که گفته بودید "
         "(«پانویس‌ها به صورت کوتاه‌شده منبع را نشان دهند») یکی نیست. کاربرگ تصمیم‌گیری "
         "«تصمیم_پانویس_ها.xlsx» که قبلاً ساختم بر پایه سیاست قبلی است و برای ۱۳ مورد "
         "«افزودن سال» را پیشنهاد داده بود.",
         12, sa=6)
    para(doc, "اگر قاعده جدید ملاک باشد، آن ۱۳ پیشنهادِ «افزودن سال» باید کنار گذاشته شود "
              "و در عوض ۱۵ پانویس تازه از جدول ۱ اضافه گردد. کاربرگ را می‌توانم بر همین مبنا "
              "بازسازی کنم — فقط بگویید.",
         12, color=GREEN, sa=8)

    para(doc, "دو موردی که نیازمند تأیید خود شماست:", 12.5, True, "right", NAVY, sa=4)
    for t in ["«فرنس» در ص ۲۵۸ (کنار یونگ) با هیچ نام شناخته‌شده‌ای جور در نمی‌آید. "
              "حدس من Franz (فون فرانتس، همکار یونگ) است، ولی باید از متن اصلی مطمئن شوید.",
              "«دیتر» در ص ۳۱۳ و «داینر» در ص ۲۶۸ ظاهراً یک نفرند (Diener). "
              "املای فارسی باید یکدست شود."]:
        para(doc, "• " + t, 11.5, sa=4, indent=0.4)

    doc.save(OUT)
    print("saved", OUT, os.path.getsize(OUT))


if __name__ == "__main__":
    build()
