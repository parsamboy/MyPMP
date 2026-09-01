# -*- coding: utf-8 -*-
"""برگهٔ اصلاح سه منبع بررسی‌شده در نسخهٔ Up-v2."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "اصلاح سه منبع (Up-v2).docx"
FA, TITR, LAT = "B Lotus", "B Titr", "Times New Roman"
NAVY=(0x1F,0x38,0x64); RED=(0xC0,0x00,0x00); GREEN=(0x1E,0x6B,0x2E)
GREY=(0x59,0x59,0x59); ORANGE=(0xB0,0x5A,0x00)


def set_rtl(p):
    b=OxmlElement("w:bidi"); b.set(qn("w:val"),"1"); p._p.get_or_add_pPr().append(b)

def sr(r,size=12,bold=False,color=None,latin=False,font=None):
    f=font or (LAT if latin else FA)
    r.font.name=f; r.font.size=Pt(size); r.font.bold=bold
    if color: r.font.color.rgb=RGBColor(*color)
    rPr=r._r.get_or_add_rPr(); rf=rPr.find(qn("w:rFonts"))
    if rf is None: rf=OxmlElement("w:rFonts"); rPr.append(rf)
    for a in ("w:cs","w:ascii","w:hAnsi"): rf.set(qn(a),f)
    s=OxmlElement("w:szCs"); s.set(qn("w:val"),str(int(size*2))); rPr.append(s)
    if not latin:
        rt=OxmlElement("w:rtl"); rt.set(qn("w:val"),"1"); rPr.append(rt)
    lg=OxmlElement("w:lang"); lg.set(qn("w:bidi"),"fa-IR"); rPr.append(lg)

def para(doc,text,size=12,bold=False,align="just",color=None,sb=0,sa=4,font=None,indent=None):
    p=doc.add_paragraph(); set_rtl(p)
    p.alignment={"just":WD_ALIGN_PARAGRAPH.JUSTIFY,"right":WD_ALIGN_PARAGRAPH.RIGHT,
                 "center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    pf=p.paragraph_format; pf.space_before=Pt(sb); pf.space_after=Pt(sa); pf.line_spacing=1.15
    if indent is not None: pf.right_indent=Cm(indent)
    if text: sr(p.add_run(text),size,bold,color,font=font)
    return p

def latin_para(doc,text,size=11,color=None,indent=0.5):
    """پاراگراف لاتین چپ‌چین برای مدخل منبع."""
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    pf=p.paragraph_format; pf.space_before=Pt(2); pf.space_after=Pt(4)
    pf.left_indent=Cm(indent); pf.line_spacing=1.1
    sr(p.add_run(text),size,False,color,latin=True)
    return p

def shade(c,hx):
    sh=OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hx)
    c._tc.get_or_add_tcPr().append(sh)

def ct(cell,text,size=10,bold=False,color=None,align="right",latin=False):
    cell.text=""; p=cell.paragraphs[0]
    if latin:
        p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    else:
        set_rtl(p)
        p.alignment={"right":WD_ALIGN_PARAGRAPH.RIGHT,"center":WD_ALIGN_PARAGRAPH.CENTER,
                     "just":WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    p.paragraph_format.space_before=Pt(1.5); p.paragraph_format.space_after=Pt(1.5)
    for i,l in enumerate(str(text).split("\n")):
        if i:
            p=cell.add_paragraph()
            if latin: p.alignment=WD_ALIGN_PARAGRAPH.LEFT
            else:
                set_rtl(p); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before=Pt(1.5); p.paragraph_format.space_after=Pt(1.5)
        sr(p.add_run(l),size,bold,color,latin=latin)

def borders(t):
    b=OxmlElement("w:tblBorders")
    for e in ("top","left","bottom","right","insideH","insideV"):
        x=OxmlElement("w:"+e); x.set(qn("w:val"),"single"); x.set(qn("w:sz"),"6")
        x.set(qn("w:color"),"808080"); b.append(x)
    t._tbl.tblPr.append(b); t._tbl.tblPr.append(OxmlElement("w:bidiVisual"))


def build():
    doc=Document(); s=doc.sections[0]
    s.page_width,s.page_height=Cm(21),Cm(29.7)
    s.top_margin=s.bottom_margin=Cm(2.0); s.left_margin=s.right_margin=Cm(1.9)
    b=OxmlElement("w:bidi"); b.set(qn("w:val"),"1"); s._sectPr.append(b)
    doc.styles["Normal"].font.name=FA; doc.styles["Normal"].font.size=Pt(12)

    para(doc,"اصلاح سه منبع",20,True,"center",NAVY,font=TITR,sa=2)
    para(doc,"بررسی نسخهٔ (Up-v2) — وضعیت منابع Zeng، Noyes و Wang & Zhao",
         12,False,"center",GREY,sa=10)

    para(doc,"پاسخ کوتاه",14,True,"right",NAVY,font=TITR,sa=4)
    para(doc,"بله، هر سه خطا در نسخهٔ Up-v2 هنوز پابرجاست. هیچ‌کدام اصلاح نشده‌اند. "
             "افزون بر این، مدخل Noyes دو بار در فهرست منابع تکرار شده و هر دو نسخه غلط است.",
         12,color=RED,sa=8)

    # ---------- ۱ Zeng ----------
    para(doc,"۱) Zeng — خطای شمارهٔ جلد",15,True,"right",ORANGE,font=TITR,sb=6,sa=4)
    para(doc,"وضعیت فعلی در فهرست منابع:",11.5,True,"right",sa=2)
    latin_para(doc,"Zeng, Y. (2011). Association of religious participation with mortality "
                   "among Chinese old adults. Research on Aging, 1: 51-83.",color=RED)
    para(doc,"سه ایراد: جلد «۱» نوشته شده که در واقع شمارهٔ نشریه است (جلد درست ۳۳ است)؛ "
             "دو نویسندهٔ دیگر (Danan Gu و Linda K. George) حذف شده‌اند؛ و قالب صفحه‌شماری "
             "با شیوه‌نامهٔ APA نمی‌خواند.",12,sa=4)
    para(doc,"جایگزین شود با:",11.5,True,"right",color=GREEN,sa=2)
    latin_para(doc,"Zeng, Y., Gu, D., & George, L. K. (2011). Association of religious "
                   "participation with mortality among Chinese old adults. Research on "
                   "Aging, 33(1), 51–83. https://doi.org/10.1177/0164027510383584",color=GREEN)
    para(doc,"همچنین ادعای متن بدنه نادرست است. متن فعلی می‌گوید این پژوهش نشان داد اعمال "
             "مذهبی «بهداشت روانی» را افزایش می‌دهد (کاهش افسردگی، افزایش عزت نفس، حمایت "
             "اجتماعی و کیفیت زندگی). اما این مطالعه هیچ‌یک از این متغیرها را نسنجیده است؛ "
             "متغیر وابستهٔ آن «مرگ‌ومیر» بوده است.",12,sb=4,sa=4)
    para(doc,"جملهٔ پیشنهادی برای متن بدنه:",11.5,True,"right",color=GREEN,sa=2)
    para(doc,"«زنگ، گو و جورج (۲۰۱۱) در پژوهشی طولی بر روی ۹٬۰۱۷ سالمند ۸۵ سال به بالا و "
             "۶٬۹۵۶ سالمند ۶۵ تا ۸۴ سال در چین نشان دادند که مشارکت منظم در فعالیت‌های مذهبی "
             "با کاهش ۲۴ درصدی خطر مرگ همراه است و این رابطه پس از تعدیل وضعیت سلامت پایه "
             "نیز با کاهش ۲۱ درصدی خطر مرگ برقرار می‌ماند.»",12,color=GREEN,sa=8,indent=0.4)

    # ---------- ۲ Noyes ----------
    doc.add_page_break()
    para(doc,"۲) Noyes — منبعِ ارجاع‌شده اصلاً وجود ندارد",15,True,"right",RED,font=TITR,sa=4)
    para(doc,"وضعیت فعلی در فهرست منابع (دو مدخل تکراری، هر دو غلط):",11.5,True,"right",sa=2)
    latin_para(doc,"Noyes, R., Stuart, S., Langbehn, D. R., & Happel, R. L. (2005). Anxiety "
                   "and fear of death in hypochondriasis. Depression and Anxiety, 21(3), 153–162.",
               color=RED)
    latin_para(doc,"Noyes, R; Stuart, S; Langbehn, D. R; & Happel, R. L. (2005). Anxiety and "
                   "fear of death in hypochondriasis. Depression and Anxiety, 21(3), 153–162.",
               color=RED)
    para(doc,"این ارجاع در سال، عنوان، نام مجله، جلد و صفحات نادرست است. جست‌وجو در PubMed "
             "نشان داد در سال ۲۰۰۵ هیچ مقاله‌ای از Noyes دربارهٔ خودبیمارانگاری در مجلهٔ "
             "Depression and Anxiety منتشر نشده است. ضمناً نام Susan L. Longley از فهرست "
             "نویسندگان جا افتاده است.",12,sa=4)
    para(doc,"هر دو مدخل حذف و یک مدخل زیر جایگزین شود:",11.5,True,"right",color=GREEN,sa=2)
    latin_para(doc,"Noyes, R., Jr., Stuart, S., Longley, S. L., Langbehn, D. R., & Happel, "
                   "R. L. (2002). Hypochondriasis and fear of death. Journal of Nervous and "
                   "Mental Disease, 190(8), 503–509. "
                   "https://doi.org/10.1097/00005053-200208000-00002",color=GREEN)
    para(doc,"ادعای متن بدنه نیز باید تعدیل شود. متن فعلی می‌گوید «در مطالعه‌ای بر روی افراد "
             "مسن»، در حالی که نمونهٔ پژوهش ۱۶۲ بیمار سرپایی عمومی بود (۴۹ نفر با تشخیص "
             "خودبیمارانگاری و ۱۱۳ نفر بدون آن)، نه سالمندان. یافتهٔ اصلی درست است و "
             "می‌توانید آن را نگه دارید.",12,sb=4,sa=4)
    para(doc,"جملهٔ پیشنهادی برای متن بدنه:",11.5,True,"right",color=GREEN,sa=2)
    para(doc,"«نویز و همکاران (۲۰۰۲) در مطالعه‌ای بر روی ۱۶۲ بیمار سرپایی گزارش کردند که "
             "ترس از مرگ همبستگی بالایی با نشانه‌های خودبیمارانگاری و اضطراب سلامتی دارد. "
             "تحلیل عاملی سه بُعد را آشکار کرد: ترس از مردن، از دست دادن معنا، و ترس از "
             "جدایی. آنان نتیجه گرفتند که ترس از مرگ جزئی جدایی‌ناپذیر از خودبیمارانگاری است.»",
         12,color=GREEN,sa=4,indent=0.4)
    para(doc,"توجه: در متن بدنه دو جا به «نویز و همکاران (۲۰۰۵)» ارجاع شده است؛ هر دو باید "
             "به (۲۰۰۲) تغییر کند.",11,color=GREY,sa=8)

    # ---------- ۳ Wang & Zhao ----------
    doc.add_page_break()
    para(doc,"۳) Wang & Zhao — منبع قابل تأیید نیست",15,True,"right",RED,font=TITR,sa=4)
    para(doc,"این ارجاع در متن بدنه آمده است:",11.5,True,"right",sa=2)
    para(doc,"«وانگ و ژائو (۲۰۲۰) ضمن پژوهشی به رابطه هوش معنوی با بهزیستی روانی و هدف در "
             "زندگی سالمندان پرداختند … ۴۰۰ سالمند از شهر شنبانگ به روش نمونه‌گیری در دسترس "
             "انتخاب شدند …»",12,color=GREY,sa=4,indent=0.4)
    para(doc,"اما هیچ مدخلی برای این منبع در فهرست منابع Up-v2 وجود ندارد. جست‌وجو در "
             "PubMed، Springer، PLOS، Google Scholar و منابع فارسی نیز هیچ مقاله‌ای با این "
             "نویسندگان، سال، موضوع و حجم نمونه نیافت.",12,sa=4)

    t=doc.add_table(rows=2,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; borders(t)
    ct(t.cell(0,0),"هشدار",size=12,bold=True,color=RED); shade(t.cell(0,0),"FBE4E4")
    ct(t.cell(1,0),"برخلاف منبع Zeng — که فقط شمارهٔ جلد جابه‌جا شده بود و با اصلاح آن "
                   "بلافاصله پیدا شد — برای این ارجاع هیچ مورد نزدیکی هم یافت نشد. "
                   "تا زمانی که اصل مقاله (DOI یا فایل PDF) در دست نباشد، نباید مشخصات "
                   "کتاب‌شناختی برای آن حدس زده شود؛ نوشتن مشخصات ساختگی در فهرست منابع "
                   "خطای علمی جدی است.",size=11.5,align="just")
    para(doc,"",size=6,sa=6)

    para(doc,"دو راه پیش رو دارید:",12,True,"right",sa=4)
    for txt in ["اصل مقاله را پیدا و ارائه کنید تا مشخصات دقیق آن استخراج شود.",
                "اگر مقاله در دسترس نیست، این ارجاع و جملهٔ وابسته به آن از بخش پیشینهٔ "
                "خارجی حذف شود. با توجه به آنکه مطالب مشابهی در همان بخش از منابع معتبرِ "
                "تأییدشده نقل شده است، حذف آن آسیبی به استدلال فصل دوم نمی‌زند."]:
        para(doc,"• "+txt,11.5,sa=5,indent=0.4)
    para(doc,"نکته: در فایل فعلی، نام «شنبانگ» احتمالاً غلط تایپی «شنیانگ» (Shenyang) است. "
             "این نکته هنگام جست‌وجوی مقاله می‌تواند کمک کند.",11,color=GREY,sb=4,sa=8)

    # ---------- جمع‌بندی ----------
    para(doc,"جمع‌بندی",15,True,"right",NAVY,font=TITR,sb=6,sa=4)
    tb=doc.add_table(rows=4,cols=4); tb.alignment=WD_TABLE_ALIGNMENT.CENTER; borders(tb)
    for i,h in enumerate(["منبع","وضعیت در Up-v2","نوع خطا","اقدام"]):
        ct(tb.cell(0,i),h,size=10.5,bold=True,align="center"); shade(tb.cell(0,i),"DCE6F1")
    rows=[("Zeng (2011)","اصلاح نشده","جلد، نویسندگان، ادعای متن","جایگزینی مدخل + بازنویسی جمله"),
          ("Noyes (2005)","اصلاح نشده + تکراری","کل ارجاع اشتباه است","حذف دو مدخل + افزودن مدخل ۲۰۰۲"),
          ("Wang & Zhao (2020)","بدون مدخل","منبع پیدا نشد","ارائهٔ اصل مقاله یا حذف ارجاع")]
    for ri,r in enumerate(rows,1):
        for ci,v in enumerate(r):
            ct(tb.cell(ri,ci),v,size=10.5,align="center" if ci else "right")
            if ri%2==1: shade(tb.cell(ri,ci),"F7F9FC")

    para(doc,"هر سه مورد الگوی یکسانی دارند: سال ارجاع غلط است و خلاصهٔ مطالعه در متن با "
             "آنچه پژوهش واقعاً سنجیده تفاوت دارد. اصلاح سال آسان است، اما ناهم‌خوانی محتوایی "
             "همان چیزی است که در جلسهٔ دفاع پرسیده می‌شود.",
         11.5,color=ORANGE,sb=8,sa=4)

    doc.save(OUT); print("saved",OUT,os.path.getsize(OUT))


if __name__=="__main__":
    build()
